"""DOCX to text converter."""

import logging
from io import BytesIO
from .base import FileConverter, ConversionResult

logger = logging.getLogger(__name__)


class DOCXConverter(FileConverter):
    """Converter for DOCX files to text."""
    
    def can_convert(self, mime_type: str, file_name: str) -> bool:
        """Check if file is a DOCX."""
        return (
            mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or
            file_name.lower().endswith(".docx")
        )
    
    async def convert(self, file_content: bytes, file_name: str, mime_type: str) -> ConversionResult:
        """Convert DOCX to text."""
        try:
            from docx import Document
            
            docx_file = BytesIO(file_content)
            doc = Document(docx_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                if table_text:
                    text_parts.append("\n[Table]\n" + "\n".join(table_text))
            
            text = "\n\n".join(text_parts)
            
            if text.strip():
                logger.info(f"✅ Converted DOCX {file_name} ({len(text)} chars)")
                return ConversionResult(
                    success=True,
                    text=text,
                    metadata={"paragraphs": len([p for p in doc.paragraphs if p.text.strip()]), "converter": "python-docx"}
                )
            else:
                return ConversionResult(
                    success=False,
                    text="",
                    error="DOCX file appears to be empty"
                )
                
        except ImportError:
            return ConversionResult(
                success=False,
                text="",
                error="python-docx library not installed"
            )
        except Exception as e:
            logger.error(f"Error converting DOCX {file_name}: {e}", exc_info=True)
            return ConversionResult(
                success=False,
                text="",
                error=f"Error extracting text from DOCX: {str(e)}"
            )

